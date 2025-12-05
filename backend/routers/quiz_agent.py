from fastapi import APIRouter, HTTPException
from models.schemas import ChatRequest, QuizResponse, QuizQuestion, QuizGenerationResponse
from services.rag_service import rag_service
import json

router = APIRouter(prefix="/api", tags=["quiz_agent"])

TEACHER_PERSONA = """你是一个新闻传播学领域的专业出题专家。
要求：
1. 根据用户指令（如“出3道关于议程设置的选择题”）生成相应数量和类型的题目。
2. JSON 结构必须如下：
{
  "questions": [
    {
      "id": 1,
      "type": "single_choice",
      "stem": "题干内容",
      "options": ["选项A", "选项B", "选项C", "选项D"],
      "answer": "A",
      "analysis": "答案解析",
      "difficulty": "medium"
    }
  ]
}
3. 如果用户没有指定数量，默认生成 1 道题。
4. **只返回 JSON 字符串**，不要包含 markdown 格式标记。
"""

STUDENT_PERSONA = """你是一个新闻传播学领域的智能辅导员。
要求：
1. 根据用户的练习需求生成试题。
2. 如果用户未指定数量，默认生成 1 道单项选择题。
3. 如果用户指定了数量（如“出3道题”），请满足其需求，但最多不超过 5 道。
4. JSON 结构必须如下：
{
  "questions": [
    {
      "id": 1,
      "type": "single_choice",
      "stem": "题干内容",
      "options": ["选项A", "选项B", "选项C", "选项D"],
      "answer": "A",
      "analysis": "答案解析（包含思考提示）",
      "difficulty": "easy"
    }
  ]
}
5. **只返回 JSON 字符串**，不要包含 markdown 格式标记。
"""

@router.post("/quiz/generate", response_model=QuizGenerationResponse)
async def generate_quiz(request: ChatRequest):
    try:
        context_str = ""
        sources = []
        
        # 1. Retrieve relevant documents (Only if use_kb is True)
        if request.use_kb:
            results = rag_service.retrieve(request.query, role=request.role, target_user_ids=request.target_user_ids)
            
            documents = results['documents'][0]
            metadatas = results['metadatas'][0]
            
            context_parts = []
            for doc, meta in zip(documents, metadatas):
                source_name = meta.get('source', 'Unknown')
                context_parts.append(f"Source ({source_name}):\n{doc}")
                if source_name not in sources:
                    sources.append(source_name)
            
            context_str = "\n\n".join(context_parts)
        
        # 2. Select Persona based on Role
        base_persona = TEACHER_PERSONA if request.role == "teacher" else STUDENT_PERSONA
        
        # 3. Construct System Prompt based on KB usage
        if request.use_kb:
            system_prompt = base_persona + "\n请基于【背景知识】和用户的【指令】，生成一组试题。内容必须基于背景知识，严谨准确。"
        else:
            system_prompt = base_persona + "\n请基于用户的【指令】和你的专业知识，生成一组试题。"
        
        # 4. Generate Answer (Force JSON)
        json_str = rag_service.generate_answer(
            query=request.query,
            context=context_str,
            history=request.history,
            system_prompt=system_prompt
        )
        
        # Clean up potential markdown code blocks if LLM adds them
        if json_str.startswith("```json"):
            json_str = json_str.replace("```json", "").replace("```", "")
        elif json_str.startswith("```"):
            json_str = json_str.replace("```", "")
            
        try:
            quiz_data = json.loads(json_str)
        except json.JSONDecodeError:
            print(f"JSON Decode Error. Raw output: {json_str}")
            # Fallback: try to find JSON object in text
            start = json_str.find("{")
            end = json_str.rfind("}") + 1
            if start != -1 and end != -1:
                try:
                    quiz_data = json.loads(json_str[start:end])
                except:
                     raise ValueError("Failed to generate valid JSON quiz.")
            else:
                raise ValueError("Failed to generate valid JSON quiz.")

        return QuizGenerationResponse(
            questions=[QuizQuestion(**q) for q in quiz_data.get("questions", [])],
            sources=sources
        )

    except Exception as e:
        print(f"Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

from fastapi.responses import StreamingResponse
from docx import Document
import io

# ... (existing code)

@router.post("/quiz/export")
async def export_quiz(quiz: QuizResponse):
    doc = Document()
    doc.add_heading(quiz.title, 0)
    
    for i, q in enumerate(quiz.questions):
        doc.add_paragraph(f"{i+1}. {q.stem}")
        for j, opt in enumerate(q.options):
            doc.add_paragraph(f"{chr(65+j)}. {opt}")
        
        doc.add_paragraph(f"正确答案: {q.answer}")
        doc.add_paragraph(f"解析: {q.analysis}")
        doc.add_paragraph("") # Spacing

    # Save to memory
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=quiz.docx"}
    )
